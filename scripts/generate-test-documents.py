from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / 'ТЕСТОВЫЕ_ДОКУМЕНТЫ_ДЛЯ_ЗАГРУЗКИ'
CONTRACT = 'ТЕСТ-701-ИЗЛК-СМР-2026'
CIPHER = 'КБ-701.26.10.01.01'
CONTRACTOR = 'ООО «ТестСтройПроект»'
INN = '7707012345'
ADDRESS = 'г. Москва, ул. Примерная, д. 7, стр. 1'


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn('w:tcW'))
    if tc_w is None:
        tc_w = OxmlElement('w:tcW')
        tc_pr.append(tc_w)
    tc_w.set(qn('w:w'), str(width_twips))
    tc_w.set(qn('w:type'), 'dxa')


def add_docx(path: Path, title: str, subtitle: str, rows: list[tuple[str, str]], body: str, parser_details: list[str] | None = None) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal = doc.styles['Normal']
    normal.font.name = 'Arial'
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)

    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run(title)
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(18)
    run.bold = True
    run.font.color.rgb = RGBColor(73, 45, 140)
    heading.paragraph_format.space_after = Pt(4)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(subtitle)
    sub_run.font.name = 'Arial'
    sub_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    sub_run.font.size = Pt(10)
    sub_run.font.color.rgb = RGBColor(90, 90, 90)
    sub.paragraph_format.space_after = Pt(14)

    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    for label, value in rows:
        cells = table.add_row().cells
        set_cell_width(cells[0], 2700)
        set_cell_width(cells[1], 6660)
        cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cells[0], 'EEEAF8')
        label_run = cells[0].paragraphs[0].add_run(label)
        label_run.bold = True
        value_run = cells[1].paragraphs[0].add_run(value)
        for run in (label_run, value_run):
            run.font.name = 'Arial'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
            run.font.size = Pt(10)

    if parser_details:
        doc.add_paragraph()
        parser_heading = doc.add_paragraph()
        parser_heading_run = parser_heading.add_run('Реквизиты для автоматического распознавания')
        parser_heading_run.bold = True
        parser_heading_run.font.color.rgb = RGBColor(73, 45, 140)
        for line in parser_details:
            parser_paragraph = doc.add_paragraph(line)
            parser_paragraph.paragraph_format.space_after = Pt(1)

    doc.add_paragraph()
    body_heading = doc.add_paragraph()
    body_heading_run = body_heading.add_run('Назначение тестового файла')
    body_heading_run.bold = True
    body_heading_run.font.color.rgb = RGBColor(73, 45, 140)
    body_paragraph = doc.add_paragraph(body)
    body_paragraph.paragraph_format.line_spacing = 1.15

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run('Тестовый комплект ИЗЛК — не является реальным коммерческим документом')
    footer_run.font.name = 'Arial'
    footer_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(120, 120, 120)

    doc.save(path)


def add_pdf(path: Path, title: str, rows: list[tuple[str, str]], body: str, parser_details: list[str] | None = None) -> None:
    font_path = Path(r'C:\Windows\Fonts\arial.ttf')
    bold_path = Path(r'C:\Windows\Fonts\arialbd.ttf')
    pdfmetrics.registerFont(TTFont('TestArial', str(font_path)))
    pdfmetrics.registerFont(TTFont('TestArialBold', str(bold_path)))
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('TestTitle', parent=styles['Title'], fontName='TestArialBold', fontSize=18, leading=22, textColor=colors.HexColor('#492D8C'), alignment=1, spaceAfter=12)
    body_style = ParagraphStyle('TestBody', parent=styles['BodyText'], fontName='TestArial', fontSize=10, leading=14, spaceAfter=8)
    label_style = ParagraphStyle('TestLabel', parent=body_style, fontName='TestArialBold')
    doc = SimpleDocTemplate(str(path), pagesize=A4, rightMargin=20 * mm, leftMargin=20 * mm, topMargin=18 * mm, bottomMargin=18 * mm)
    story = [Paragraph(title, title_style), Paragraph('Тестовый документ для загрузки в ИЗЛК', body_style)]
    table_data = [[Paragraph(label, label_style), Paragraph(value, body_style)] for label, value in rows]
    table = Table(table_data, colWidths=[48 * mm, 122 * mm])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#EEEAF8')),
        ('GRID', (0, 0), (-1, -1), 0.35, colors.HexColor('#C8C1E1')),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('LEFTPADDING', (0, 0), (-1, -1), 7),
        ('RIGHTPADDING', (0, 0), (-1, -1), 7),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
    ]))
    story.extend([table, Spacer(1, 12)])
    if parser_details:
        story.append(Paragraph('<b>Реквизиты для автоматического распознавания</b>', label_style))
        story.extend(Paragraph(line, body_style) for line in parser_details)
    story.extend([Paragraph('<b>Назначение тестового файла</b>', label_style), Paragraph(body, body_style)])
    doc.build(story)


def main() -> None:
    folders = [
        '01_Парсер_договора',
        '02_Хронология_договора',
        '03_Проект',
        '04_Исполнительная_документация',
        '05_Прочее',
        '06_Для_очереди_импорта',
    ]
    for folder in folders:
        (OUT / folder).mkdir(parents=True, exist_ok=True)

    common = [
        ('Номер договора', CONTRACT),
        ('Шифр объекта', CIPHER),
        ('Контрагент', CONTRACTOR),
        ('ИНН', INN),
        ('Адрес объекта', ADDRESS),
    ]
    contract_rows = common + [('Дата договора', '05.08.2026'), ('Сумма договора', '15 780 000,00 руб.'), ('Валюта', 'RUB')]
    parser_details = [
        f'Договор № {CONTRACT} от 05.08.2026',
        f'Контрагент: {CONTRACTOR}',
        f'ИНН: {INN}',
        f'Шифр объекта: {CIPHER}',
        f'Адрес объекта: {ADDRESS}',
        'Сумма договора: 15 780 000,00 руб.',
    ]
    add_docx(OUT / '01_Парсер_договора' / f'Договор_{CONTRACT}.docx', 'ДОГОВОР ПОДРЯДА', f'№ {CONTRACT}', contract_rows, 'Используйте этот файл в разделе «Договоры → Загрузить договор». Система должна найти номер, дату, сумму, контрагента, ИНН, шифр и адрес.', parser_details)
    add_pdf(OUT / '01_Парсер_договора' / f'Договор_{CONTRACT}.pdf', 'ДОГОВОР ПОДРЯДА', contract_rows, 'PDF содержит текстовый слой и подходит для проверки парсера PDF.', parser_details)

    add_docx(OUT / '02_Хронология_договора' / f'Дополнительное_соглашение_01_{CONTRACT}.docx', 'ДОПОЛНИТЕЛЬНОЕ СОГЛАШЕНИЕ № 1', f'к договору {CONTRACT}', common + [('Дата', '10.08.2026'), ('Изменение', 'Уточнение состава работ')], 'После создания договора загрузите файл в карточку договора и выберите тип «Дополнительное соглашение».')
    add_docx(OUT / '02_Хронология_договора' / f'Коммерческое_предложение_{CONTRACT}.docx', 'КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ', f'по объекту {CIPHER}', common + [('Дата', '04.08.2026'), ('Стоимость', '15 780 000,00 руб.')], 'Загружайте как «Коммерческое предложение».')
    add_pdf(OUT / '02_Хронология_договора' / f'Счет_01_{CONTRACT}.pdf', 'СЧЕТ НА ОПЛАТУ № 01', common + [('Дата счета', '06.08.2026'), ('Сумма к оплате', '3 156 000,00 руб.'), ('Срок оплаты', '15.08.2026')], 'Загружайте как «Счет на оплату». Бухгалтерские статусы остаются в 1С.')

    add_pdf(OUT / '03_Проект' / f'КМ_{CIPHER}.pdf', 'ПРОЕКТНАЯ ДОКУМЕНТАЦИЯ КМ', common + [('Раздел', 'КМ'), ('Стадия', 'Рабочая документация')], 'Загрузите в карточку договора как «Проект PDF» и привяжите к разделу КМ.')
    add_pdf(OUT / '03_Проект' / f'КЖ_{CIPHER}.pdf', 'ПРОЕКТНАЯ ДОКУМЕНТАЦИЯ КЖ', common + [('Раздел', 'КЖ'), ('Стадия', 'Рабочая документация')], 'Загрузите в карточку договора как «Проект PDF» и привяжите к разделу КЖ.')
    add_pdf(OUT / '03_Проект' / f'АР_{CIPHER}.pdf', 'ПРОЕКТНАЯ ДОКУМЕНТАЦИЯ АР', common + [('Раздел', 'АР'), ('Стадия', 'Рабочая документация')], 'Загрузите в карточку договора как «Проект PDF» и привяжите к разделу АР.')

    exec_base = OUT / '04_Исполнительная_документация'
    add_docx(exec_base / f'ОЖР_{CONTRACT}.docx', 'ОБЩИЙ ЖУРНАЛ РАБОТ', f'по договору {CONTRACT}', common + [('Период', 'Август 2026'), ('Статус', 'Ведение начато')], 'При загрузке выберите исполнительный документ и раздел «ОЖР».')
    add_docx(exec_base / f'Паспорт_на_каркас_{CONTRACT}.docx', 'ПАСПОРТ НА КАРКАС', f'по объекту {CIPHER}', common + [('Марка металла', 'С245'), ('Статус', 'Готов к выдаче')], 'При загрузке выберите исполнительный документ и раздел «Паспорт на каркас».')
    add_docx(exec_base / f'Приказ_о_назначении_{CONTRACT}.docx', 'ПРИКАЗ О НАЗНАЧЕНИИ ОТВЕТСТВЕННЫХ ЛИЦ', f'по договору {CONTRACT}', common + [('Дата приказа', '05.08.2026'), ('Ответственный', 'Иванов И.И.')], 'При загрузке выберите исполнительный документ и раздел «Акты и приказы».')
    add_pdf(exec_base / f'Сертификат_на_металл_{CONTRACT}.pdf', 'СЕРТИФИКАТ СООТВЕТСТВИЯ', common + [('Материал', 'Прокат стальной'), ('Номер сертификата', 'ТЕСТ-2026-701')], 'При загрузке выберите исполнительный документ и раздел «Сертификаты».')
    add_pdf(exec_base / f'Схема_расположения_{CONTRACT}.pdf', 'СХЕМА РАСПОЛОЖЕНИЯ КОЛОНН', common + [('Раздел', 'КМ'), ('Листы', '1-3')], 'При загрузке выберите исполнительный документ и раздел «Схемы».')

    (OUT / '05_Прочее' / f'Примечание_технадзора_{CONTRACT}.txt').write_text(
        f'ЗАМЕЧАНИЕ ТЕХНАДЗОРА\nДоговор: {CONTRACT}\nШифр: {CIPHER}\nДата: 08.08.2026\nТекст: Тестовое замечание для проверки раздела «Прочее».\n',
        encoding='utf-8',
    )
    (OUT / '05_Прочее' / f'Реестр_для_импорта_{CONTRACT}.csv').write_text(
        'Поле;Значение\nНомер договора;' + CONTRACT + '\nШифр;' + CIPHER + '\nКонтрагент;' + CONTRACTOR + '\nИНН;' + INN + '\nСумма;15780000\n',
        encoding='utf-8-sig',
    )
    (OUT / '06_Для_очереди_импорта' / f'Договор_{CONTRACT}_для_очереди.txt').write_text(
        f'Договор № {CONTRACT}\nШифр: {CIPHER}\nКонтрагент: {CONTRACTOR}\nИНН: {INN}\nДата: 05.08.2026\nСумма: 15 780 000 руб.\n',
        encoding='utf-8',
    )
    (OUT / 'README_КАК_ТЕСТИРОВАТЬ.txt').write_text(
        'ТЕСТОВЫЙ КОМПЛЕКТ ИЗЛК\n\n'
        '1. В папке 01_Парсер_договора загрузите DOCX или PDF через «Договоры → Загрузить договор».\n'
        '2. Создайте карточку договора из распознанных полей.\n'
        '3. Файлы из 02 и 03 загружайте в карточку созданного договора, выбирая тип файла.\n'
        '4. В 04 лежат файлы для разных разделов исполнительной документации. На странице ИД нажмите «Добавить файл в раздел».\n'
        '5. Файл из 06 скопируйте в inbox и проверьте «Очередь импорта». Точную копию положите повторно: она не должна появиться второй раз.\n'
        '6. Все файлы тестовые, их можно удалять из карточки после проверки.\n',
        encoding='utf-8',
    )


if __name__ == '__main__':
    main()
